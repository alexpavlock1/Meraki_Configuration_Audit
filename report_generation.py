import json
import shutil
import csv
from prettytable import PrettyTable
from docx import Document
import logging
import os
import zipfile
import meraki.aio
from docx.shared import RGBColor


logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

def load_json(filepath):
    with open(filepath, 'r') as file:
        return json.load(file)

def load_prettytable_from_file(filepath):
    logging.debug(f"Loading PrettyTable from file: {filepath}")
    try:
        with open(filepath, 'r') as file:
            table_string = file.read()
        
        lines = table_string.split('\n')
        if len(lines) < 3:
            logging.warning(f"File {filepath} has insufficient data. Creating empty table.")
            table = PrettyTable()
            return table
            
        field_names = [field.strip() for field in lines[1].split('|')[1:-1]]
        rows = [line.split('|')[1:-1] for line in lines[3:-1] if line.strip()]
        
        table = PrettyTable()
        table.field_names = field_names
        for row in rows:
            if row:  # Skip empty rows
                table.add_row([item.strip() for item in row])
        
        return table
    except FileNotFoundError:
        logging.warning(f"File {filepath} not found. Creating empty table.")
        table = PrettyTable()
        table.field_names = ["No Data Available"]
        table.add_row(["Table data not generated during compliance check"])
        return table
    except Exception as e:
        logging.error(f"Error loading table from {filepath}: {e}")
        table = PrettyTable()
        table.field_names = ["Error"]
        table.add_row([f"Error loading table: {e}"])
        return table

def add_prettytable_to_doc(doc, table, title, description=None):
    logging.debug("Adding PrettyTable to document")
    doc.add_heading(title, level=2)
    
    if description:
        doc.add_paragraph(description)
    
    rows = len(table._rows) + 1
    cols = len(table.field_names)
    table_word = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    
    for i, field in enumerate(table.field_names):
        cell = table_word.cell(0, i)
        cell.text = field
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, row in enumerate(table._rows):
        for j, value in enumerate(row):
            cell = table_word.cell(i + 1, j)
            cell.text = str(value).strip()

def zip_files(directory, zip_filename):
    with zipfile.ZipFile(zip_filename, 'w') as zipf:
        for root, _, files in os.walk(directory):
            for file in files:
                zipf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), directory))
    logging.debug(f"Created zip file {zip_filename}")

def json_to_csv(json_data, csv_filepath):
    with open(csv_filepath, 'w', newline='') as file:
        writer = csv.writer(file)
        if isinstance(json_data, list):
            if len(json_data) > 0 and isinstance(json_data[0], dict):
                writer.writerow(json_data[0].keys())
                for row in json_data:
                    writer.writerow(row.values())
            else:
                writer.writerow(['Value'])
                for row in json_data:
                    writer.writerow([row])
        elif isinstance(json_data, dict):
            writer.writerow(json_data.keys())
            writer.writerow(json_data.values())
        else:
            writer.writerow(['Value'])
            writer.writerow([json_data])
    logging.debug(f"Converted JSON data to CSV at {csv_filepath}")

async def generate_report(report_directory):
    logging.debug("Generating report")
    summary_table = load_prettytable_from_file('summary_table.txt')
    non_compliant_networks_table = load_prettytable_from_file('non_compliant_networks_table.txt')
    verdict_table = load_prettytable_from_file('verdict_table.txt')
    network_verdict_table = load_prettytable_from_file('network_verdict_table.txt')

    doc = Document('CiscoTemplate.docx')
    doc.add_heading('Compliance Report', level=1)
    overall_status = "All configurations are compliant"
    for row in summary_table._rows:
        if row[3].strip().lower() != "compliant":
            overall_status = "Not all configurations are compliant"
            break

    overall_summary_table = PrettyTable()
    overall_summary_table.field_names = ["Overall Status"]
    overall_summary_table.add_row([overall_status])

    add_prettytable_to_doc(doc, overall_summary_table, "Overall Summary")

    session_data = load_json('session_data.json')
    org_names = session_data['org_names']
    golden_images = session_data['golden_images']
    api_key = session_data['api_key']
    template_ids = session_data.get('template_ids', [])  # Get template_ids from session_data

    async with meraki.aio.AsyncDashboardAPI(api_key, output_log=False) as dashboard:
        golden_networks_table = PrettyTable()
        golden_networks_table.field_names = ["Organization", "Golden Network"]

        for org_id, golden_network_id in golden_images.items():
            org_name = org_names[org_id]
            
            # Check if the golden network is a template
            is_golden_template = golden_network_id in template_ids
            
            try:
                if is_golden_template:
                    # For templates, fetch template info
                    templates = await dashboard.organizations.getOrganizationConfigTemplates(org_id)
                    golden_template = next((t for t in templates if t['id'] == golden_network_id), None)
                    if golden_template:
                        golden_network_name = golden_template['name'] + ' (Template)'
                    else:
                        golden_network_name = "Template_" + golden_network_id
                else:
                    # For regular networks, fetch network info
                    golden_network = await dashboard.networks.getNetwork(golden_network_id)
                    golden_network_name = golden_network['name']
            except Exception as e:
                logging.error(f"Error fetching golden network/template {golden_network_id}: {e}")
                # Fallback: use a generic name
                if is_golden_template:
                    golden_network_name = "Template_" + golden_network_id
                else:
                    golden_network_name = "Network_" + golden_network_id
            
            golden_networks_table.add_row([org_name, golden_network_name])

    add_prettytable_to_doc(doc, golden_networks_table, "Golden Networks")

    add_prettytable_to_doc(doc, network_verdict_table, "Network Verdicts")
    add_prettytable_to_doc(doc, verdict_table, "Verdict Details")

    # Load AI analysis results
    try:
        with open('anomalies.json', 'r') as file:
            anomalies = json.load(file)
        logging.debug(f"Loaded anomalies data")
    except Exception as e:
        logging.warning(f"Failed to load anomalies: {e}")
        anomalies = {}
        
    try:
        with open('recommendations.json', 'r') as file:
            recommendations = json.load(file)
        logging.debug(f"Loaded recommendations data")
    except Exception as e:
        logging.warning(f"Failed to load recommendations: {e}")
        recommendations = {}
    try:
        with open('check_name_mapping.json', 'r') as file:
            check_name_mapping = json.load(file)
        logging.debug(f"Loaded check name mapping")
    except Exception as e:
        logging.warning(f"Failed to load check name mapping: {e}")
        check_name_mapping = {}
    logging.info(f"check_name_mapping loaded successfully with {len(check_name_mapping)} entries")
    # Create a reverse mapping from endpoint names to display names
    endpoint_to_display_name = {}
    for api_name, display_name in check_name_mapping.items():
        # Extract the endpoint name without the prefixes
        endpoint_name = api_name.replace('getNetwork', '').replace('getOrganization', '').replace('Appliance', '').replace('Switch', '').lower()
        endpoint_to_display_name[endpoint_name] = display_name
    #logging.info(f"Created endpoint_to_display_name mapping with {len(endpoint_to_display_name)} entries")
    #logging.info(f"Sample mappings: {list(endpoint_to_display_name.items())[:5]}")
    # Add anomalies section if data is available
    if anomalies:
        doc.add_heading("Security Anomalies", level=2)
        doc.add_paragraph("The following networks show unusual security configurations that may require attention.")
        
        # Create a table for anomalies
        anomalies_table = doc.add_table(rows=1, cols=4, style='Table Grid')
        header_cells = anomalies_table.rows[0].cells
        header_cells[0].text = "Network"
        header_cells[1].text = "Risk Level"
        header_cells[2].text = "Score"
        header_cells[3].text = "Security Issues"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add anomaly data to table
        for network_id, anomaly in anomalies.items():
            row = anomalies_table.add_row()
            cells = row.cells
            
            # Network name
            cells[0].text = anomaly.get('network_name', network_id)
            
            # Risk level with color
            risk_level = anomaly.get('risk_level', 'Unknown')
            cells[1].text = risk_level
            
            # Make risk level cells colored based on severity
            if risk_level == "Critical":
                for paragraph in cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                        run.font.bold = True
            elif risk_level == "High":
                for paragraph in cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 102, 0)  # Orange
                        run.font.bold = True
            elif risk_level == "Medium":
                for paragraph in cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 204, 0)  # Yellow
            
            # Anomaly score
            cells[2].text = str(anomaly.get('anomaly_score', '-'))
            
            # Security issues with bullet points
            issues_text = ""
            for issue in anomaly.get('security_issues', []):
                issues_text += f"• {issue.get('severity', '')} ({issue.get('category', '')}): {issue.get('description', '')}\n"
            
            if not issues_text:
                issues_text = "No specific issues identified"
                
            cells[3].text = issues_text
    else:
        # Add a note in the report that no anomalies were detected
        doc.add_heading("Security Anomalies", level=2)
        doc.add_paragraph("No security anomalies were detected or anomaly data was not available.")

    # Add recommendations section if data is available
    if recommendations:
        doc.add_heading("Compliance Recommendations", level=2)
        doc.add_paragraph("The following recommendations will help bring non-compliant networks into compliance.")
        
        # Process each network's recommendations
        for network_id, network_recs in recommendations.items():
            #logging.info(f"Network {network_id} has {len(network_recs)} recommendations")
            # Try to find network name from anomalies or use ID if not found
            network_name = network_id
            for row in network_verdict_table._rows:
                # Find the right row that has the network ID
                found = False
                for org_id, networks in session_data.get('compliance_networks', {}).items():
                    if network_id in networks:
                        # This org has this network
                        org_name = org_names.get(org_id, "Unknown Org")
                        if row[0] == org_name:  # Match organization
                            network_name = row[1]  # Use network name from table
                            found = True
                            break
                if found:
                    break
                    
            # Also check if name is in anomalies
            for net_id, anomaly in anomalies.items():
                if net_id == network_id and 'network_name' in anomaly:
                    network_name = anomaly['network_name']
                    break

            
            # Add network heading
            doc.add_heading(f"Recommendations for {network_name}", level=3)
            
            # Process each recommendation for this network
            for rec in network_recs:
                # Create a recommendation section
                rec_title = rec.get('title', 'Configuration Issue')
                #logging.info(f"Processing recommendation with title: '{rec_title}'")
                # Try to find a friendly name for the title using our mapping
                friendly_title = rec_title
                found_match = False
                for endpoint_name, display_name in endpoint_to_display_name.items():
                    # Extract the base name (remove " Configuration" if present)
                    base_title = rec_title.replace(" Configuration", "").lower()
                    if base_title == endpoint_name.lower():
                        #logging.info(f"Found match: '{rec_title}' -> '{display_name}'")
                        friendly_title = display_name
                        found_match = True
                        break

            if not found_match:
                logging.warning(f"No match found for '{rec_title}' in {len(endpoint_to_display_name)} mappings")
                
            category = rec.get('category', 'General')
            doc.add_paragraph(f"{friendly_title} ({category})", style='Heading 4')
                
            # Process each action
            for action in rec.get('actions', []):
                action_type = action.get('action_type', 'Action')
                description = action.get('description', '')
                    
                # Add action heading with color based on type
                p = doc.add_paragraph()
                run = p.add_run(f"{action_type}: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                p.add_run(description)
                
                # Add details as bullet points if any
                if 'details' in action and action['details']:
                    for detail in action['details']:
                        p = doc.add_paragraph(detail, style='List Bullet')
                
            # Add separator
            doc.add_paragraph("─" * 40)
    else:
        # No recommendations
        if network_verdict_table._rows and any("Non-Compliant" in row[-1] for row in network_verdict_table._rows):
            doc.add_heading("Compliance Recommendations", level=2)
            doc.add_paragraph("Recommendations data could not be generated for non-compliant networks.")

    report_path = os.path.join(report_directory, 'Meraki_Health_Check.docx')
    doc.save(report_path)
    logging.debug(f"Report generated and saved as '{report_path}'")

    for org_id, org_name in org_names.items():
        org_directory = f"Org_{org_name.replace(' ', '_')}"
        if os.path.exists(org_directory):
            logging.debug(f"Processing organization directory: {org_directory}")
            for root, _, files in os.walk(org_directory):
                for file in files:
                    if file.endswith('.json'):
                        src_path = os.path.join(root, file)
                        dest_path = os.path.join(report_directory, os.path.relpath(src_path, org_directory))
                        os.makedirs(os.path.dirname(dest_path), exist_ok=True)
                        shutil.copy(src_path, dest_path)
                        logging.debug(f"Copied JSON data file '{src_path}' to '{dest_path}'")
                        
                        json_data = load_json(src_path)
                        csv_filepath = dest_path.replace('.json', '.csv')
                        json_to_csv(json_data, csv_filepath)
                        logging.debug(f"Converted JSON data to CSV at '{csv_filepath}'")

async def main():
    try:
        await generate_report()
    except Exception as e:
        logging.error(f"Report generation failed: {e}")

if __name__ == '__main__':
    asyncio.run(main())